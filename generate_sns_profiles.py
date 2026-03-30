"""
Aiden SNSプロフィール完全設計書
Word（.docx）出力
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── カラー定数 ──────────────────────────────────────────────────
C_GOLD        = RGBColor(0xC9, 0xA0, 0x2A)
C_DEEPBLUE    = RGBColor(0x0D, 0x2A, 0x4E)
C_JADE        = RGBColor(0x00, 0x7A, 0x5E)
C_PURPLE      = RGBColor(0x4A, 0x14, 0x8C)
C_WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
C_DARK        = RGBColor(0x1A, 0x1A, 0x2E)
C_LIGHT_GOLD  = RGBColor(0xFF, 0xF0, 0xCC)
C_LIGHT_BLUE  = RGBColor(0xE8, 0xF4, 0xFF)
C_LIGHT_GREEN = RGBColor(0xE8, 0xF5, 0xE9)
C_LIGHT_PURPLE= RGBColor(0xF3, 0xE5, 0xF5)
C_LIGHT_GRAY  = RGBColor(0xF5, 0xF5, 0xF5)
C_TIKTOK      = RGBColor(0xFE, 0x2C, 0x55)
C_INSTA       = RGBColor(0x83, 0x3A, 0xB4)
C_LINE        = RGBColor(0x06, 0xC7, 0x55)
C_X           = RGBColor(0x00, 0x00, 0x00)
C_ORANGE      = RGBColor(0xE6, 0x51, 0x00)

# ── ヘルパー ────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), str(rgb))
    tcPr.append(shd)


def run_font(run, size=10.5, bold=False, color=None, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = "Meiryo UI"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Meiryo UI")
    if color:
        run.font.color.rgb = color


def add_para(doc_or_cell, text, align=WD_ALIGN_PARAGRAPH.LEFT,
             size=10.5, bold=False, color=None, italic=False,
             space_before=0, space_after=4, indent=0):
    if hasattr(doc_or_cell, 'add_paragraph'):
        p = doc_or_cell.add_paragraph()
    else:
        p = doc_or_cell.paragraphs[0] if not doc_or_cell.paragraphs[0].text else doc_or_cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.left_indent  = Pt(indent)
    if text:
        run = p.add_run(text)
        run_font(run, size=size, bold=bold, color=color, italic=italic)
    return p


def title_banner(doc, text, bg: RGBColor, fg: RGBColor = C_WHITE, size=15):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Pt(10)
    run = p.add_run(text)
    run_font(run, size=size, bold=True, color=fg)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def section_banner(doc, text, bg: RGBColor, fg: RGBColor = C_WHITE, size=12):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(10)
    run = p.add_run(text)
    run_font(run, size=size, bold=True, color=fg)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def sub_banner(doc, text, bg: RGBColor, fg: RGBColor = C_WHITE, size=10.5):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Pt(10)
    run = p.add_run(text)
    run_font(run, size=size, bold=True, color=fg)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def profile_table(doc, rows, col_w=(3.8, 12.6), label_bg=None, val_bg=None,
                  label_fg=C_DARK, note_rows=None):
    """ラベル｜値 の2列テーブル。note_rows はセルを薄色にする行インデックスリスト"""
    note_rows = note_rows or []
    n = len(rows)
    tbl = doc.add_table(rows=n, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(rows):
        row = tbl.rows[i]
        row.cells[0].width = Cm(col_w[0])
        row.cells[1].width = Cm(col_w[1])

        # ラベル
        lb_bg = label_bg if label_bg else C_LIGHT_GRAY
        set_cell_bg(row.cells[0], lb_bg)
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p0.paragraph_format.space_before = Pt(4)
        p0.paragraph_format.space_after  = Pt(4)
        p0.paragraph_format.left_indent  = Pt(6)
        r0 = p0.add_run(label)
        run_font(r0, size=10, bold=True, color=label_fg)

        # 値
        is_note = i in note_rows
        v_bg = RGBColor(0xFA, 0xFA, 0xEE) if is_note else (val_bg if val_bg else C_WHITE)
        set_cell_bg(row.cells[1], v_bg)

        if isinstance(value, list):
            first = True
            for line in value:
                if first:
                    p1 = row.cells[1].paragraphs[0]
                    first = False
                else:
                    p1 = row.cells[1].add_paragraph()
                p1.paragraph_format.space_before = Pt(2)
                p1.paragraph_format.space_after  = Pt(2)
                p1.paragraph_format.left_indent  = Pt(8)
                r1 = p1.add_run(line)
                run_font(r1, size=10, italic=is_note,
                         color=RGBColor(0x55, 0x55, 0x55) if is_note else C_DARK)
        else:
            p1 = row.cells[1].paragraphs[0]
            p1.paragraph_format.space_before = Pt(4)
            p1.paragraph_format.space_after  = Pt(4)
            p1.paragraph_format.left_indent  = Pt(8)
            r1 = p1.add_run(str(value))
            run_font(r1, size=10.5 if not is_note else 10,
                     bold=not is_note,
                     italic=is_note,
                     color=RGBColor(0x55, 0x55, 0x55) if is_note else C_DARK)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def tip_box(doc, title, lines, bg: RGBColor, title_bg: RGBColor = None):
    """ヒント・解説ボックス"""
    tb_bg = title_bg if title_bg else bg
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # タイトル行
    ct = tbl.cell(0, 0)
    set_cell_bg(ct, tb_bg)
    pt = ct.paragraphs[0]
    pt.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pt.paragraph_format.space_before = Pt(3)
    pt.paragraph_format.space_after  = Pt(3)
    pt.paragraph_format.left_indent  = Pt(10)
    rt = pt.add_run(title)
    run_font(rt, size=10.5, bold=True, color=C_WHITE)

    # 本文行
    cb = tbl.cell(1, 0)
    set_cell_bg(cb, bg)
    first = True
    for line in lines:
        if first:
            pb = cb.paragraphs[0]
            first = False
        else:
            pb = cb.add_paragraph()
        pb.paragraph_format.space_before = Pt(2)
        pb.paragraph_format.space_after  = Pt(2)
        pb.paragraph_format.left_indent  = Pt(12)
        if line:
            rb = pb.add_run(line)
            run_font(rb, size=10)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def variations_table(doc, headers, rows_data, header_bg: RGBColor,
                     col_widths, alt_bg: RGBColor = None):
    """バリエーション比較テーブル"""
    tbl = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    alt = alt_bg if alt_bg else C_LIGHT_GRAY

    for c_i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = tbl.cell(0, c_i)
        cell.width = Cm(w)
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(h)
        run_font(r, size=10, bold=True, color=C_WHITE)

    for r_i, row_vals in enumerate(rows_data, 1):
        row = tbl.rows[r_i]
        bg = alt if r_i % 2 == 1 else C_WHITE
        for c_i, val in enumerate(row_vals):
            cell = row.cells[c_i]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Pt(6)
            if isinstance(val, list):
                first = True
                for line in val:
                    if first:
                        r = p.add_run(line)
                        run_font(r, size=10, bold=(c_i == 0))
                        first = False
                    else:
                        px = cell.add_paragraph()
                        px.paragraph_format.space_before = Pt(0)
                        px.paragraph_format.space_after  = Pt(2)
                        px.paragraph_format.left_indent  = Pt(6)
                        rx = px.add_run(line)
                        run_font(rx, size=10)
            else:
                r = p.add_run(str(val))
                run_font(r, size=10, bold=(c_i == 0))

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ════════════════════════════════════════════════════════════════
# コンテンツ定義
# ════════════════════════════════════════════════════════════════

TIKTOK_PROFILES = [
    {
        "type": "推奨①（メイン）",
        "account_name_en": "Aiden | AI Business",
        "account_name_jp": "Aiden｜AI自動化の専門家",
        "username": "@aiden_ai_business",
        "bio": "✦ AIで業務を自動化するプロ✦\n中小企業の社長の時間を取り戻す🐉\n無料相談→プロフのLINEへ",
        "bio_count": "60文字",
        "link_desc": "🐉 無料AI診断はこちら → LINE登録",
        "point": "龍神シンボル（🐉）で差別化。「時間を取り戻す」が経営者の痛みに直撃。短くて記憶に残る構成。",
    },
    {
        "type": "推奨②（共感型）",
        "account_name_en": "Aiden | Work Smart",
        "account_name_jp": "Aiden｜残業ゼロへの道",
        "username": "@aiden_worksmart",
        "bio": "残業・手作業・繰り返し仕事\nぜんぶAIに任せませんか？🌿\n無料相談はプロフのLINEから✨",
        "bio_count": "57文字",
        "link_desc": "✨ 無料でAI活用診断 → LINEで相談",
        "point": "ターゲットの「日常の悩み」を3連で列挙。翡翠エネルギー（🌿）を使用。問いかけ形式で共感を誘発。",
    },
    {
        "type": "推奨③（権威型）",
        "account_name_en": "Aiden | AI Consultant",
        "account_name_jp": "Aiden｜AI導入コンサル",
        "username": "@aiden_consultant",
        "bio": "50社超のAI導入実績✦\nあなたの会社も自動化できます\n📩 無料相談 → プロフLINE",
        "bio_count": "51文字",
        "link_desc": "📩 50社導入実績のAI相談 → LINE",
        "point": "実績数字を冒頭に配置（権威性）。「あなたの会社も」でターゲットの壁を下げる。",
    },
]

INSTA_PROFILES = [
    {
        "type": "推奨①（メイン・バランス型）",
        "account_name": "Aiden｜AI業務自動化の専門家",
        "username": "@aiden_ai_official",
        "bio": (
            "✦ AIで業務を自動化するプロ ✦\n"
            "🐉 中小企業・個人事業主の\n"
            "　　時間と利益を最大化する\n\n"
            "📊 導入実績50社以上\n"
            "⏰ 平均2.5時間/日の業務削減\n\n"
            "👇 無料AI診断はLINEから"
        ),
        "bio_count": "105文字",
        "link_desc": "🐉 無料AI活用診断（30分）→ LINE登録はこちら",
        "highlights": ["AI活用術", "導入事例", "料金プラン", "無料相談", "会社情報"],
        "point": "数字で権威性確立。龍神シンボルで視覚的差別化。改行を活かした読みやすい構成。",
    },
    {
        "type": "推奨②（スピリチュアル×ビジネス融合型）",
        "account_name": "Aiden｜龍神に愛されるAI活用",
        "username": "@aiden_ryujin_ai",
        "bio": (
            "✨ 龍神の加護を受けたAIアドバイザー\n\n"
            "🌿 あなたの会社に眠る\n"
            "　　可能性をAIで解き放つ\n\n"
            "💰 業務自動化で月30時間を取り戻す\n"
            "🔗 無料相談はプロフLINEへ"
        ),
        "bio_count": "112文字",
        "link_desc": "✨ 龍神エネルギー×AI → 無料相談LINE",
        "highlights": ["龍神×AI", "自動化事例", "無料診断", "サービス", "お客様の声"],
        "point": "スピリチュアル×ビジネスの唯一無二ポジション。Z世代・スピリチュアル好き経営者に刺さる。",
    },
    {
        "type": "推奨③（シンプル権威型）",
        "account_name": "Aiden / AI Business Advisor",
        "username": "@aiden_biz_ai",
        "bio": (
            "AI業務自動化アドバイザー\n\n"
            "✅ チャットボット開発\n"
            "✅ 業務プロセス自動化\n"
            "✅ AI導入コンサルティング\n\n"
            "📩 無料相談 → 下のLINEから"
        ),
        "bio_count": "88文字",
        "link_desc": "📩 サービス詳細・無料相談 → LINE",
        "highlights": ["サービス一覧", "料金プラン", "導入事例", "AI活用Tips", "無料相談"],
        "point": "チェックリスト形式でサービスが一目瞭然。ビジネス色が強く、真剣な経営者に刺さる。",
    },
]

INSTA_HIGHLIGHT_DETAILS = [
    ("AI活用術",  "🤖", "AIの基本・使い方・ビジネス活用のノウハウ投稿をまとめる\n定期更新してフォロワーの「辞書」になる存在を目指す"),
    ("導入事例",  "📈", "お客様のビフォーアフター・成果数字・業種別事例\n社会的証明の宝庫として最も重要なハイライト"),
    ("料金プラン","💰", "スタータープラン〜エンタープライズの説明\n「高そう」という先入観を事前に解消する役割"),
    ("無料相談",  "📩", "無料相談の流れ・Q&A・LINE登録方法\nCTA（行動促進）の中心として常に最前列に配置"),
    ("会社情報",  "🏢", "Aiden株式会社の理念・メンバー・実績\n権威性と信頼感を積み上げるブランドの土台"),
]

LINE_PROFILES = [
    {
        "type": "推奨①（温かみ重視）",
        "account_name": "Aiden｜AI業務自動化相談室",
        "status": "✦ 今月の無料相談：残り3枠 ✦",
        "greeting": (
            "友達追加ありがとうございます！🐉\n"
            "Aidenです。AIで業務を自動化するプロとして、\n"
            "中小企業・個人事業主のみなさまを\n"
            "全力でサポートしています✨\n\n"
            "━━━━━━━━━━━━━\n"
            "✅ Aidenができること\n"
            "━━━━━━━━━━━━━\n"
            "🤖 AIチャットボット開発（¥50,000〜）\n"
            "⚙️ 業務プロセス自動化（¥150,000〜）\n"
            "📊 AI導入コンサルティング（¥300,000〜）\n\n"
            "━━━━━━━━━━━━━\n"
            "🎁 今なら無料でプレゼント\n"
            "━━━━━━━━━━━━━\n"
            "「あなたの会社で自動化できる業務診断」\n"
            "30分の無料相談でお届けします！\n\n"
            "👇 ご希望の方は「診断希望」と送ってください\n\n"
            "一緒に、あなたの会社を\n"
            "もっと楽に・もっと豊かにしていきましょう🌿\n\n"
            "Aiden（エイデン）"
        ),
        "point": "希少性（残り3枠）でステータスに緊張感。あいさつは「価値提供→行動促進」の構成。",
    },
    {
        "type": "推奨②（シンプル・行動直結型）",
        "account_name": "Aiden AI｜無料相談はここから",
        "status": "💬 メッセージください → 24時間以内に返信します",
        "greeting": (
            "こんにちは！Aidenです🐉\n"
            "AIで業務自動化のプロをしています。\n\n"
            "友達追加してくださった方に、\n"
            "今だけ特典をプレゼントしています👇\n\n"
            "🎁【無料特典】\n"
            "「業務自動化チェックシート（PDF）」\n"
            "あなたの会社で今すぐ自動化できる\n"
            "業務を診断できるシートです！\n\n"
            "受け取り方法：\n"
            "このトークに「チェックシート」と\n"
            "送ってください！すぐお届けします✨\n\n"
            "また、無料AI相談も受付中です。\n"
            "「相談したい」とメッセージいただければ\n"
            "日程を調整します😊\n\n"
            "— Aiden"
        ),
        "point": "無料特典で即行動を促す。「チェックシート」という具体的キーワードが行動ハードルを下げる。",
    },
]

X_PROFILES = [
    {
        "type": "推奨①（権威×共感バランス型）",
        "account_name": "Aiden｜AI業務自動化の専門家",
        "username": "@aiden_ai_pro",
        "bio": (
            "AIで業務を自動化するプロ🐉｜中小企業・個人事業主の"
            "「時間と利益」を最大化｜導入実績50社以上｜"
            "チャットボット/業務自動化/AIコンサル｜"
            "無料相談→プロフのLINEから📩"
        ),
        "bio_count": "109文字",
        "pin_tweet": (
            "【保存版】中小企業がAIで真っ先に自動化すべき業務5選🤖\n\n"
            "1️⃣ 問い合わせ対応（LINEチャットボット）\n"
            "2️⃣ データ入力・集計（RPAツール）\n"
            "3️⃣ 請求書作成・送付（経理自動化）\n"
            "4️⃣ SNS投稿・メール返信（AI文章生成）\n"
            "5️⃣ スケジュール調整（AIアシスタント）\n\n"
            "これだけで月30時間が浮きます。\n"
            "年間360時間＝丸15日分の自由時間が生まれます。\n\n"
            "まず1つだけ始めてみてください。\n"
            "どこから手をつければいいか悩む方は\n"
            "プロフのLINEから無料相談どうぞ👇\n\n"
            "#AI活用 #業務自動化 #中小企業 #DX #時間節約"
        ),
        "point": "保存促進コンテンツをピン留め。リスト形式で最後まで読まれる。相談誘導がシームレス。",
    },
    {
        "type": "推奨②（個性・スピリチュアル型）",
        "account_name": "Aiden🐉｜龍神×AI",
        "username": "@aiden_ryujin",
        "bio": (
            "龍神の加護を受けたAIアドバイザー🐉✨｜"
            "中小企業の社長の「時間・お金・自由」をAIで取り戻す｜"
            "弁財天エネルギーで言葉を届ける｜"
            "無料AI診断→プロフLINE"
        ),
        "bio_count": "97文字",
        "pin_tweet": (
            "社長に質問です。\n\n"
            "あなたが1日で一番時間を使っている作業、\n"
            "何ですか？\n\n"
            "実はその作業、\n"
            "高確率でAIに任せられます。\n\n"
            "データ入力・メール対応・請求書作成・\n"
            "SNS投稿・スケジュール調整...\n\n"
            "「それ全部、AIでできますよ」\n\n"
            "信じてもらえないかもしれないけど、\n"
            "本当の話です。\n\n"
            "試しに一度、無料で話してみてください。\n"
            "プロフのLINEから30分だけ。\n\n"
            "後悔しない確率、かなり高いです🐉\n\n"
            "#AI活用 #中小企業社長 #業務効率化 #龍神"
        ),
        "point": "スピリチュアル×ビジネスの独自ポジション。ピン留めは「会話形式」でミラーニューロンを刺激。",
    },
    {
        "type": "推奨③（ビジネス特化型）",
        "account_name": "Aiden | AI Consultant",
        "username": "@aiden_ai_consult",
        "bio": (
            "AIソリューション専門家｜"
            "チャットボット開発・業務自動化・AI導入支援｜"
            "導入50社/平均ROI300%｜"
            "毎日AI活用Tipsを発信｜"
            "無料相談はプロフリンクのLINEから📩"
        ),
        "bio_count": "102文字",
        "pin_tweet": (
            "AI導入で失敗する会社と成功する会社\n"
            "たった1つの違い。\n\n"
            "失敗する会社：「なんとなく導入してみよう」\n"
            "成功する会社：「月10時間のこの作業をゼロにする」\n\n"
            "目的が明確なだけで、\n"
            "成功率が10倍変わります。\n\n"
            "うちのクライアントには必ず最初に聞きます。\n"
            "「何を解決したいですか？」\n\n"
            "答えられたら、あとは任せてください。\n"
            "最短2週間で動くものを作ります。\n\n"
            "詳しくはプロフのLINEから👇\n"
            "#AI導入 #DX #業務効率化 #中小企業 #チャットボット"
        ),
        "point": "ロジカルで信頼性重視。「ROI300%」の数字で即座に権威確立。ビジネス系フォロワーに強い。",
    },
]

COMMON_CONCEPTS = [
    ("ブランドカラー",      "深藍（#0D2A4E）× 皇帝ゴールド（#C9A02A）× 翡翠緑（#007A5E）"),
    ("シンボルマーク",      "🐉 龍神シンボル（差別化・上昇・金運の象徴）"),
    ("サブシンボル",        "🌿 翡翠緑（弁財天・成長・安心）/ ✦ 星型（品格・特別感）"),
    ("アカウント共通思想",  "「AIは怖くない・難しくない・高くない」という固定概念を壊す存在"),
    ("一貫したCTA",         "全プラットフォーム共通：「無料相談はプロフのLINEから」"),
    ("言葉のトーン",        "専門的だけど親しみやすい。難しい言葉を使わない。数字で話す。"),
    ("禁止ワード",          "「絶対」「確実に稼げる」「全員成功」→ 誇大表現はすべてNG"),
    ("投稿テーマの柱",      "①AI活用ノウハウ ②導入事例・数字 ③経営者への共感 ④CTA"),
]

# ════════════════════════════════════════════════════════════════
# ドキュメント構築
# ════════════════════════════════════════════════════════════════

def build():
    doc = Document()

    # ── ページ設定 ──────────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(1.8)
    sec.right_margin  = Cm(1.8)
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = "Meiryo UI"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Meiryo UI")

    # ════════════════════════════════════════════════════════════
    # 表紙
    # ════════════════════════════════════════════════════════════
    cov = doc.add_table(rows=1, cols=1)
    cov.alignment = WD_TABLE_ALIGNMENT.CENTER
    cc = cov.cell(0, 0)
    set_cell_bg(cc, C_DEEPBLUE)

    cover_lines = [
        ("✦ AIDEN SNS PROFILE DESIGN BOOK ✦",       10,  True,  C_GOLD),
        ("",                                           5,  False, C_WHITE),
        ("Aiden（エイデン）",                         26,  True,  C_GOLD),
        ("SNSプロフィール完全設計書",                 18,  True,  C_WHITE),
        ("",                                           4,  False, C_WHITE),
        ("TikTok  /  Instagram  /  LINE  /  X",       12, False,  RGBColor(0xAA, 0xCC, 0xFF)),
        ("",                                           3,  False, C_WHITE),
        ("— 龍神・弁財天エネルギー × AI業務自動化 —", 10, False,  C_LIGHT_GOLD),
        ("",                                           4,  False, C_WHITE),
        ("Aiden株式会社  /  2026年 最新版",           10, False,  RGBColor(0xAA, 0xCC, 0xFF)),
    ]
    first = True
    for text, size, bold, color in cover_lines:
        p = cc.paragraphs[0] if first else cc.add_paragraph()
        first = False
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        if text:
            r = p.add_run(text)
            r.font.name = "Meiryo UI"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Meiryo UI")
            r.font.size   = Pt(size)
            r.font.bold   = bold
            r.font.color.rgb = color
    for _ in range(3):
        cc.add_paragraph()

    doc.add_paragraph()
    doc.add_page_break()

    # ── 目次 ────────────────────────────────────────────────────
    title_banner(doc, "📋 目次 / CONTENTS", C_DEEPBLUE, size=14)

    toc = [
        ("第1章", "共通コンセプト設計",      "全SNS共通のブランドルール・カラー・言語設計"),
        ("第2章", "TikTok プロフィール",     "アカウント名・ユーザーネーム・自己紹介文3パターン"),
        ("第3章", "Instagram プロフィール",  "プロフィール3パターン・ハイライト5個の設計"),
        ("第4章", "LINE公式アカウント",      "あいさつメッセージ2パターン・ステータス設計"),
        ("第5章", "X（Twitter）プロフィール","自己紹介文3パターン・ピン留めツイート全文"),
        ("付録",  "文字数確認シート",        "各SNSの文字数制限と実際の文字数チェック表"),
    ]

    toc_tbl = doc.add_table(rows=len(toc), cols=3)
    toc_tbl.style = "Table Grid"
    toc_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_i, (chap, ttl, desc) in enumerate(toc):
        bg = C_LIGHT_BLUE if r_i % 2 == 0 else C_WHITE
        for c_i, (val, w) in enumerate(zip([chap, ttl, desc], [1.8, 4.8, 10.6])):
            cell = toc_tbl.cell(r_i, c_i)
            cell.width = Cm(w)
            set_cell_bg(cell, C_LIGHT_GOLD if c_i == 0 else bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after  = Pt(5)
            p.paragraph_format.left_indent  = Pt(0) if c_i == 0 else Pt(8)
            r = p.add_run(val)
            run_font(r, size=10, bold=(c_i < 2),
                     color=C_DEEPBLUE if c_i == 0 else C_DARK)

    doc.add_paragraph()
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 第1章 共通コンセプト
    # ════════════════════════════════════════════════════════════
    title_banner(doc, "第1章　全SNS共通コンセプト設計", C_DEEPBLUE, size=14)

    section_banner(doc, "1-1　ブランドルール一覧", C_DARK)
    profile_table(doc, COMMON_CONCEPTS,
                  col_w=(3.8, 12.6),
                  label_bg=RGBColor(0xBB, 0xDE, 0xFF),
                  val_bg=C_LIGHT_BLUE)

    tip_box(doc, "✦ 全SNS共通の「最重要ルール」",
            ["① 全プラットフォームで同じ配色・シンボルを使い「見た瞬間にAidenとわかる」状態を作る",
             "② すべての自己紹介の最後は「無料相談はプロフのLINEから」で統一する",
             "③ ユーザーネームは覚えやすいシンプルなものを使用し、変更しない",
             "④ アイコン画像は全SNS共通（深藍背景＋ゴールド「A」ロゴ＋Aidenの顔）",
             "⑤ 更新・投稿を継続することで「生きているアカウント」として認識させる"],
            bg=C_LIGHT_BLUE, title_bg=C_DEEPBLUE)

    section_banner(doc, "1-2　プロフィール設計の心理学原則", C_DARK)
    profile_table(doc, [
        ("ハロー効果",    "最初の0.05秒で「信頼できる人」を伝える。アイコン・配色・肩書きが命"),
        ("権威性確立",    "実績数字（50社・平均2.5時間削減）を必ず入れる。具体性が信頼を生む"),
        ("社会的証明",    "フォロワー数・導入社数・お客様の声を可視化する"),
        ("希少性",        "ステータス・あいさつ文に「今月残り〇枠」など限定感を入れる"),
        ("返報性",        "プロフィールに無料特典の存在を匂わせて「得られそう」感を出す"),
        ("一貫性",        "「無料相談だけでもOK」という小さなYESから始める導線を設計する"),
    ], label_bg=RGBColor(0xC5, 0xCA, 0xFF), val_bg=C_LIGHT_PURPLE)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 第2章 TikTok
    # ════════════════════════════════════════════════════════════
    title_banner(doc, "第2章　TikTok プロフィール設計", C_TIKTOK, size=14)

    section_banner(doc, "2-1　TikTokプロフィール仕様（文字数制限）",
                   RGBColor(0x88, 0x00, 0x00))
    profile_table(doc, [
        ("アカウント名",   "最大30文字（日本語は全角1文字＝2文字換算の場合あり）"),
        ("ユーザーネーム", "最大24文字。英数字・アンダースコア・ピリオドのみ使用可"),
        ("自己紹介文",     "最大80文字。改行可能（3〜4行推奨）"),
        ("リンク",         "プロアカウントのみ設定可能。説明文は任意（30文字以内推奨）"),
        ("重要ポイント",   "【TikTok特有の注意点】\nビジネスアカウントに切り替えること。\n「音楽」タブが無効化されるが、アナリティクスとリンク設置が可能になる。\nカテゴリを「ビジネスサービス」に設定するとリーチが最適化される。"),
    ], label_bg=RGBColor(0xFF, 0xCC, 0xCC), val_bg=RGBColor(0xFF, 0xF5, 0xF5), note_rows=[4])

    for i, p in enumerate(TIKTOK_PROFILES):
        section_banner(doc, f"2-{i+2}　{p['type']}", C_TIKTOK)
        profile_table(doc, [
            ("アカウント名（英語）",  p["account_name_en"]),
            ("アカウント名（日本語）",p["account_name_jp"]),
            ("ユーザーネーム",        p["username"]),
            ("自己紹介文",            p["bio"].split("\n")),
            ("文字数",                p["bio_count"] + "（制限80文字以内）"),
            ("リンク説明文",          p["link_desc"]),
            ("設計ポイント",          p["point"]),
        ], label_bg=RGBColor(0xFF, 0xCC, 0xCC), val_bg=C_WHITE, note_rows=[6])

    tip_box(doc, "✦ TikTok プロフィール 最終チェックリスト",
            ["□ アカウント名に「AIビジネス・自動化・コンサル」などのキーワードが入っているか",
             "□ ユーザーネームが短くて覚えやすいか（15文字以内が理想）",
             "□ 自己紹介文の最初の1行で「誰のためのアカウントか」が伝わるか",
             "□ 🐉または🌿のシンボルが入っているか（差別化と視覚的記憶のため）",
             "□ 「無料相談→プロフのLINE」の導線が明確か",
             "□ プロアカウントに設定済みでリンクが設置できているか"],
            bg=RGBColor(0xFF, 0xEB, 0xEE), title_bg=C_TIKTOK)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 第3章 Instagram
    # ════════════════════════════════════════════════════════════
    title_banner(doc, "第3章　Instagram プロフィール設計", C_INSTA, size=14)

    section_banner(doc, "3-1　Instagramプロフィール仕様（文字数制限）",
                   RGBColor(0x44, 0x00, 0x66))
    profile_table(doc, [
        ("アカウント名",    "最大30文字。検索に使われるので重要キーワードを入れる"),
        ("ユーザーネーム",  "最大30文字。英数字・アンダースコア・ピリオドのみ"),
        ("自己紹介文",      "最大150文字。改行・絵文字可能（5〜7行推奨）"),
        ("リンク",          "1つのみ設定可能（リンクツリー等を活用してLINEへ誘導）"),
        ("ハイライト名",    "最大15文字。アイコン＋色でブランドとの統一感を出す"),
        ("重要ポイント",    "【Instagram特有の注意点】\nプロアカウント（クリエイター or ビジネス）に切り替えること。\nカテゴリを「コンサルタント」または「IT企業」に設定推奨。\n自己紹介の「名前欄」は検索で使われるため最重要キーワードを入れる。"),
    ], label_bg=RGBColor(0xCC, 0xBB, 0xFF), val_bg=C_LIGHT_PURPLE, note_rows=[5])

    for i, p in enumerate(INSTA_PROFILES):
        section_banner(doc, f"3-{i+2}　{p['type']}", C_INSTA)
        profile_table(doc, [
            ("アカウント名",  p["account_name"]),
            ("ユーザーネーム",p["username"]),
            ("自己紹介文",    p["bio"].split("\n")),
            ("文字数",        p["bio_count"] + "（制限150文字以内）"),
            ("リンク説明文",  p["link_desc"]),
            ("設計ポイント",  p["point"]),
        ], label_bg=RGBColor(0xCC, 0xBB, 0xFF), val_bg=C_WHITE, note_rows=[5])

    section_banner(doc, "3-5　ハイライト設計（全プロフィール共通推奨）", C_INSTA)
    variations_table(doc,
        headers=["ハイライト名", "絵文字", "収録コンテンツ・目的"],
        rows_data=[(h, e, d) for h, e, d in INSTA_HIGHLIGHT_DETAILS],
        header_bg=C_INSTA,
        col_widths=[3.0, 1.5, 12.7],
        alt_bg=C_LIGHT_PURPLE)

    tip_box(doc, "✦ Instagram プロフィール 最終チェックリスト",
            ["□ 名前欄に「AI・業務自動化・中小企業」などのSEOキーワードが入っているか",
             "□ 自己紹介の1行目で「誰が・何をしているか」が瞬時に伝わるか",
             "□ 実績数字（50社・2.5時間削減など）が入っているか",
             "□ 龍神シンボル（🐉）または翡翠（🌿）が入っているか",
             "□ 「LINE」への誘導が最終行に必ず入っているか",
             "□ ハイライトのアイコンがブランドカラー（深藍・ゴールド）で統一されているか",
             "□ プロアカウントに切り替え済みで、インサイトが確認できるか"],
            bg=C_LIGHT_PURPLE, title_bg=C_INSTA)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 第4章 LINE公式アカウント
    # ════════════════════════════════════════════════════════════
    title_banner(doc, "第4章　LINE公式アカウント設計", C_LINE,
                 fg=RGBColor(0x00, 0x55, 0x20), size=14)

    section_banner(doc, "4-1　LINE公式アカウント仕様", RGBColor(0x02, 0x75, 0x3A))
    profile_table(doc, [
        ("アカウント名",        "最大20文字。検索されることを想定してキーワードを入れる"),
        ("ステータスメッセージ","最大120文字。アカウントの下に表示。希少性・緊急性を入れると効果的"),
        ("あいさつメッセージ",  "友達追加時に自動送信。最大500文字（複数メッセージ可）"),
        ("プロフィール画像",    "深藍背景＋ゴールドの「A」ロゴ推奨。正方形・高解像度で"),
        ("カバー画像",          "1080×878px推奨。「AI業務自動化相談室」のテキスト＋Aidenの立ち姿"),
        ("重要ポイント",        "【LINE特有の設計ポイント】\nあいさつメッセージは「価値提供→信頼構築→行動促進」の順番を守る。\n無料特典（チェックシートPDF等）をプレゼントするキーワード設定で\nユーザーの行動を促進する。リッチメニューで主要メニューを常時表示する。"),
    ], label_bg=RGBColor(0x99, 0xFF, 0xCC), val_bg=C_LIGHT_GREEN, note_rows=[5])

    for i, p in enumerate(LINE_PROFILES):
        section_banner(doc, f"4-{i+2}　{p['type']}", RGBColor(0x02, 0x9A, 0x50))
        profile_table(doc, [
            ("アカウント名",        p["account_name"]),
            ("ステータスメッセージ",p["status"]),
            ("あいさつメッセージ",  p["greeting"].split("\n")),
            ("設計ポイント",        p["point"]),
        ], label_bg=RGBColor(0x99, 0xFF, 0xCC), val_bg=C_WHITE, note_rows=[3])

    section_banner(doc, "4-4　リッチメニュー設計（推奨構成）", RGBColor(0x02, 0x75, 0x3A))
    variations_table(doc,
        headers=["メニュー位置", "ボタン名", "リンク先・機能", "カラー"],
        rows_data=[
            ("左上（大）",   "📩 無料相談",   "LINEトークへの返信促進\n「相談したい」と送ってね",   "ゴールド背景"),
            ("右上",         "🤖 AI活用Tips", "公式サイト or Instagramへ\nAI活用記事・動画へ誘導",   "深藍背景"),
            ("左下",         "💰 料金プラン", "Webサイトの料金ページへ\n「思ったより安い」体験を提供","翡翠緑背景"),
            ("中央下",       "📊 導入事例",   "事例紹介ページへ\n社会的証明を確認してもらう",       "深藍背景"),
            ("右下",         "📞 会社情報",   "Aiden株式会社の概要ページへ\n権威性・信頼性を確認",   "グレー背景"),
        ],
        header_bg=RGBColor(0x02, 0x75, 0x3A),
        col_widths=[2.8, 3.0, 7.0, 3.4],
        alt_bg=C_LIGHT_GREEN)

    tip_box(doc, "✦ LINE公式アカウント 最終チェックリスト",
            ["□ アカウント名に検索キーワード（AI・相談・自動化）が入っているか",
             "□ ステータスメッセージに希少性・緊急感（残り〇枠）が入っているか",
             "□ あいさつメッセージは価値提供→行動促進の順番になっているか",
             "□ 無料特典（PDF・診断等）のキーワード設定が済んでいるか",
             "□ リッチメニューが設置されていて「無料相談」が最も目立つ位置か",
             "□ プロフィール画像・カバー画像がブランドカラーで統一されているか"],
            bg=C_LIGHT_GREEN, title_bg=RGBColor(0x02, 0x75, 0x3A))

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 第5章 X（Twitter）
    # ════════════════════════════════════════════════════════════
    title_banner(doc, "第5章　X（Twitter）プロフィール設計", C_X, size=14)

    section_banner(doc, "5-1　X（Twitter）プロフィール仕様", RGBColor(0x33, 0x33, 0x33))
    profile_table(doc, [
        ("アカウント名",   "最大50文字。絵文字可能。検索に影響するためキーワード重要"),
        ("ユーザーネーム", "最大15文字（@なし）。英数字・アンダースコアのみ"),
        ("自己紹介文",     "最大160文字。改行・絵文字・URLなど自由度が高い"),
        ("ピン留めツイート","最も読まれるコンテンツ。保存・シェアされやすい内容を設置"),
        ("ヘッダー画像",   "1500×500px推奨。深藍グラデ＋ゴールドのキャッチコピー＋Aidenの姿"),
        ("重要ポイント",   "【X（Twitter）特有の設計ポイント】\n自己紹介は「｜」で区切ると読みやすく、情報量を多く詰め込める。\nピン留めは「保存したくなる価値ある情報」に設定することで\n新規プロフィール訪問者の滞在時間とフォロー率が上がる。"),
    ], label_bg=RGBColor(0xCC, 0xCC, 0xCC), val_bg=C_LIGHT_GRAY, note_rows=[5])

    for i, p in enumerate(X_PROFILES):
        section_banner(doc, f"5-{i+2}　{p['type']}", RGBColor(0x22, 0x22, 0x22))
        profile_table(doc, [
            ("アカウント名",    p["account_name"]),
            ("ユーザーネーム",  p["username"]),
            ("自己紹介文",      p["bio"]),
            ("文字数",          p["bio_count"] + "（制限160文字以内）"),
            ("ピン留めツイート",p["pin_tweet"].split("\n")),
            ("設計ポイント",    p["point"]),
        ], label_bg=RGBColor(0xCC, 0xCC, 0xCC), val_bg=C_WHITE, note_rows=[5])

    tip_box(doc, "✦ X（Twitter）プロフィール 最終チェックリスト",
            ["□ アカウント名に「AI｜自動化｜コンサル」などのキーワードが入っているか",
             "□ ユーザーネームは15文字以内で覚えやすいか",
             "□ 自己紹介は「｜」区切りで情報を最大限詰め込んでいるか",
             "□ 実績数字（50社・ROI300%など）が入っているか",
             "□ ピン留めツイートは「保存したくなる価値ある情報」になっているか",
             "□ ピン留めの最後に「プロフのLINEから無料相談」の誘導があるか",
             "□ ヘッダー画像がブランドカラー（深藍×ゴールド）で設計されているか"],
            bg=C_LIGHT_GRAY, title_bg=RGBColor(0x22, 0x22, 0x22))

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 付録：文字数確認シート
    # ════════════════════════════════════════════════════════════
    title_banner(doc, "付録　文字数制限・確認シート", C_DEEPBLUE, size=14)

    section_banner(doc, "各SNS 文字数制限 早見表", C_DARK)
    variations_table(doc,
        headers=["SNS", "項目", "文字数制限", "Aiden推奨文字数", "ポイント"],
        rows_data=[
            ("TikTok",      "アカウント名",    "30文字",  "15〜20文字", "キーワード＋ブランド名"),
            ("TikTok",      "ユーザーネーム",  "24文字",  "10〜15文字", "覚えやすい英数字"),
            ("TikTok",      "自己紹介文",      "80文字",  "55〜65文字", "3行で完結させる"),
            ("TikTok",      "リンク説明文",    "30文字",  "20〜25文字", "CTA＋シンボルで引き付ける"),
            ("Instagram",   "アカウント名",    "30文字",  "15〜25文字", "検索キーワードを入れる"),
            ("Instagram",   "ユーザーネーム",  "30文字",  "12〜18文字", "他SNSと統一推奨"),
            ("Instagram",   "自己紹介文",      "150文字", "100〜120文字","5〜6行で改行を活用"),
            ("Instagram",   "ハイライト名",    "15文字",  "4〜6文字",  "短くて直感的に"),
            ("LINE公式",    "アカウント名",    "20文字",  "12〜18文字", "検索されることを意識"),
            ("LINE公式",    "ステータス",      "120文字", "40〜60文字", "希少性・緊急性を入れる"),
            ("LINE公式",    "あいさつメッセージ","500文字","300〜400文字","価値提供→CTA構成"),
            ("X（Twitter）","アカウント名",    "50文字",  "15〜25文字", "｜区切りで情報量UP"),
            ("X（Twitter）","ユーザーネーム",  "15文字",  "10〜13文字", "他SNSと統一推奨"),
            ("X（Twitter）","自己紹介文",      "160文字", "100〜120文字","｜区切り形式が読みやすい"),
        ],
        header_bg=C_DEEPBLUE,
        col_widths=[2.8, 3.6, 2.8, 3.6, 4.4],
        alt_bg=C_LIGHT_BLUE)

    section_banner(doc, "SNS別 推奨ユーザーネーム候補一覧", C_DARK)
    variations_table(doc,
        headers=["パターン", "TikTok", "Instagram", "X（Twitter）", "統一性"],
        rows_data=[
            ("メイン推奨",   "@aiden_ai_business", "@aiden_ai_official",  "@aiden_ai_pro",      "高（aidenで統一）"),
            ("スピリチュアル系","@aiden_ryujin",   "@aiden_ryujin_ai",    "@aiden_ryujin",      "最高（完全統一）"),
            ("ビジネス特化", "@aiden_consultant",  "@aiden_biz_ai",       "@aiden_ai_consult",  "中（aidenで統一）"),
            ("シンプル",     "@aiden_official",    "@aiden_official",     "@aiden_official",    "最高（完全統一）"),
        ],
        header_bg=C_DEEPBLUE,
        col_widths=[3.0, 4.0, 4.5, 4.2, 1.5],
        alt_bg=C_LIGHT_BLUE)

    tip_box(doc, "✦ 最終推奨：全SNS統一ユーザーネーム",
            ["推奨①（スピリチュアル×ビジネス路線）：@aiden_ryujin",
             "   → 「龍神×AI」という唯一無二のポジションを全SNSで一貫して確立できる",
             "   → 覚えやすく、検索しやすく、差別化が最大化される",
             "",
             "推奨②（シンプル・ビジネス路線）：@aiden_official",
             "   → 全SNSで完全統一。ブランド名そのものなので信頼感が高い",
             "   → 将来的に企業アカウントに移行する際もそのまま使いやすい",
             "",
             "※ どちらのユーザーネームも取得可能かどうか、各SNSで確認してから使用すること"],
            bg=C_LIGHT_GOLD, title_bg=C_GOLD, title_bg2=C_GOLD)

    # ── フッター ───────────────────────────────────────────────
    doc.add_paragraph()
    tip_box(doc, "✦ 本設計書について",
            ["本書はAiden株式会社のブランドアバター「Aiden（エイデン）」の",
             "SNSプロフィール完全設計書です。",
             "本書のプロフィール文はすべてコピー＆ペーストしてそのまま使用できます。",
             "",
             "制作：Aiden株式会社 マーケティング部門　/ 2026年3月　VERSION 1.0"],
            bg=C_LIGHT_BLUE, title_bg=C_DEEPBLUE)

    out = "/home/eiko/aiden/Aiden_SNS_Profile_Design_Book.docx"
    doc.save(out)
    print(f"✅ 完成: {out}")


if __name__ == "__main__":
    build()
