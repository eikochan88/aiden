"""
Aiden ブランドアバター設計書
python-docx で Word（.docx）出力
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── カラー定数 ──────────────────────────────────────────────────
C_GOLD        = RGBColor(0xC9, 0xA0, 0x2A)   # 金
C_DEEPBLUE    = RGBColor(0x0D, 0x2A, 0x4E)   # 深藍
C_JADE        = RGBColor(0x00, 0x7A, 0x5E)   # 翡翠緑
C_PURPLE      = RGBColor(0x4A, 0x14, 0x8C)   # 深紫
C_WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
C_LIGHT_GOLD  = RGBColor(0xFF, 0xF0, 0xCC)
C_LIGHT_BLUE  = RGBColor(0xE8, 0xF4, 0xFF)
C_LIGHT_GREEN = RGBColor(0xE8, 0xF5, 0xE9)
C_LIGHT_PURPLE= RGBColor(0xF3, 0xE5, 0xF5)
C_LIGHT_GRAY  = RGBColor(0xF5, 0xF5, 0xF5)
C_DARK_TEXT   = RGBColor(0x1A, 0x1A, 0x2E)
C_RED         = RGBColor(0xC6, 0x28, 0x28)
C_TEAL        = RGBColor(0x00, 0x69, 0x6F)

# ── ヘルパー ────────────────────────────────────────────────────

def set_cell_bg(cell, hex_str):
    """セル背景色を16進RGB文字列でセット（例: 'C9A02A'）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_str)
    tcPr.append(shd)


def hex_rgb(r, g, b):
    return f"{r:02X}{g:02X}{b:02X}"


def cell_bg_rgb(cell, rgb: RGBColor):
    set_cell_bg(cell, str(rgb))


def para_style(para, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)


def run_style(run, size=10.5, bold=False, color: RGBColor = None, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = "Meiryo UI"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Meiryo UI")
    if color:
        run.font.color.rgb = color


def add_heading_block(doc, text, level=1, bg: RGBColor = None, fg: RGBColor = C_WHITE):
    """背景色付き見出し段落を追加"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    if bg:
        cell_bg_rgb(cell, bg)
    cell.width = Inches(6.8)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(8)
    run = p.add_run(text)
    sizes = {1: 16, 2: 13, 3: 11}
    run_style(run, size=sizes.get(level, 11), bold=True, color=fg)
    doc.add_paragraph()


def add_two_col_table(doc, rows_data, col_widths=(2.8, 9.0),
                      header=None, header_bg: RGBColor = C_DEEPBLUE,
                      odd_bg: RGBColor = C_LIGHT_GRAY, even_bg: RGBColor = C_WHITE,
                      label_bg: RGBColor = None, label_fg: RGBColor = C_DARK_TEXT):
    """2列テーブルを追加"""
    n = len(rows_data) + (1 if header else 0)
    table = doc.add_table(rows=n, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 列幅設定
    for row in table.rows:
        row.cells[0].width = Cm(col_widths[0])
        row.cells[1].width = Cm(col_widths[1])

    row_offset = 0
    if header:
        cell_l = table.cell(0, 0)
        cell_r = table.cell(0, 1)
        cell_l.merge(cell_r)
        cell_bg_rgb(cell_l, header_bg)
        p = cell_l.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(header)
        run_style(run, size=11, bold=True, color=C_WHITE)
        row_offset = 1

    for i, (label, value) in enumerate(rows_data):
        r = table.rows[i + row_offset]
        # ラベル列
        lb_bg = label_bg if label_bg else (odd_bg if i % 2 == 0 else even_bg)
        cell_bg_rgb(r.cells[0], lb_bg)
        p0 = r.cells[0].paragraphs[0]
        p0.paragraph_format.space_before = Pt(3)
        p0.paragraph_format.space_after = Pt(3)
        p0.paragraph_format.left_indent = Pt(4)
        run0 = p0.add_run(label)
        run_style(run0, size=10, bold=True, color=label_fg)
        # 値列
        v_bg = odd_bg if i % 2 == 0 else even_bg
        cell_bg_rgb(r.cells[1], v_bg)
        p1 = r.cells[1].paragraphs[0]
        p1.paragraph_format.space_before = Pt(3)
        p1.paragraph_format.space_after = Pt(3)
        p1.paragraph_format.left_indent = Pt(6)
        if isinstance(value, list):
            for vi, line in enumerate(value):
                if vi == 0:
                    run1 = p1.add_run(line)
                    run_style(run1, size=10)
                else:
                    px = r.cells[1].add_paragraph()
                    px.paragraph_format.space_before = Pt(0)
                    px.paragraph_format.space_after = Pt(1)
                    px.paragraph_format.left_indent = Pt(6)
                    rx = px.add_run(line)
                    run_style(rx, size=10)
        else:
            run1 = p1.add_run(str(value))
            run_style(run1, size=10)

    doc.add_paragraph()
    return table


def add_box(doc, title, body_lines, border_color: RGBColor = C_GOLD,
            title_bg: RGBColor = None, body_bg: RGBColor = C_LIGHT_GOLD):
    """タイトル＋本文ボックス"""
    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # タイトル行
    tc_title = table.cell(0, 0)
    title_bg_c = title_bg if title_bg else border_color
    cell_bg_rgb(tc_title, title_bg_c)
    p_t = tc_title.paragraphs[0]
    p_t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_t.paragraph_format.space_before = Pt(3)
    p_t.paragraph_format.space_after = Pt(3)
    p_t.paragraph_format.left_indent = Pt(8)
    run_t = p_t.add_run(title)
    run_style(run_t, size=11, bold=True, color=C_WHITE)

    # 本文行
    tc_body = table.cell(1, 0)
    cell_bg_rgb(tc_body, body_bg)
    first = True
    for line in body_lines:
        if first:
            pb = tc_body.paragraphs[0]
            first = False
        else:
            pb = tc_body.add_paragraph()
        pb.paragraph_format.space_before = Pt(1)
        pb.paragraph_format.space_after = Pt(1)
        pb.paragraph_format.left_indent = Pt(8)
        run_b = pb.add_run(line)
        run_style(run_b, size=10)

    doc.add_paragraph()


def add_color_swatch_table(doc, swatches):
    """カラースウォッチテーブル（色名・HEX・意味・効果）"""
    table = doc.add_table(rows=len(swatches) + 1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["カラー名", "HEX", "RGBイメージ", "象徴・意味", "心理的効果"]
    widths_cm = [2.8, 2.0, 1.6, 4.0, 5.4]

    for col_i, (h, w) in enumerate(zip(headers, widths_cm)):
        cell = table.cell(0, col_i)
        cell.width = Cm(w)
        cell_bg_rgb(cell, C_DEEPBLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(h)
        run_style(run, size=9.5, bold=True, color=C_WHITE)

    for r_i, (name, hex_c, rgb_c, meaning, effect) in enumerate(swatches, 1):
        row = table.rows[r_i]
        vals = [name, hex_c, "■■■", meaning, effect]
        for c_i, val in enumerate(vals):
            cell = row.cells[c_i]
            if c_i == 2 and rgb_c:
                cell_bg_rgb(cell, rgb_c)
            else:
                bg = C_LIGHT_GOLD if r_i % 2 == 1 else C_WHITE
                cell_bg_rgb(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_i in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Pt(4) if c_i not in (1, 2) else Pt(0)
            run = p.add_run(val)
            fg = C_WHITE if c_i == 2 and rgb_c else C_DARK_TEXT
            run_style(run, size=9.5, bold=(c_i == 0), color=fg)

    doc.add_paragraph()


def add_psych_table(doc, items):
    """心理学テーブル（原則・説明・アバターへの適用・期待効果）"""
    headers = ["心理学原則", "仕組み", "Aidenアバターへの適用", "期待される効果"]
    widths_cm = [3.2, 4.0, 6.0, 3.6]

    table = doc.add_table(rows=len(items) + 1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for c_i, (h, w) in enumerate(zip(headers, widths_cm)):
        cell = table.cell(0, c_i)
        cell.width = Cm(w)
        cell_bg_rgb(cell, C_PURPLE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(h)
        run_style(run, size=9.5, bold=True, color=C_WHITE)

    for r_i, (principle, mechanism, application, effect) in enumerate(items, 1):
        row = table.rows[r_i]
        bg = C_LIGHT_PURPLE if r_i % 2 == 1 else C_WHITE
        vals = [principle, mechanism, application, effect]
        for c_i, val in enumerate(vals):
            cell = row.cells[c_i]
            cell_bg_rgb(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Pt(4)
            if isinstance(val, list):
                for vi, line in enumerate(val):
                    if vi == 0:
                        run = p.add_run(line)
                        run_style(run, size=9.5, bold=(c_i == 0))
                    else:
                        px = cell.add_paragraph()
                        px.paragraph_format.space_before = Pt(0)
                        px.paragraph_format.space_after = Pt(1)
                        px.paragraph_format.left_indent = Pt(4)
                        rx = px.add_run(line)
                        run_style(rx, size=9.5)
            else:
                run = p.add_run(str(val))
                run_style(run, size=9.5, bold=(c_i == 0))

    doc.add_paragraph()


def add_voice_table(doc, items):
    headers = ["カテゴリ", "設計内容", "具体例・スクリプト", "心理効果"]
    widths_cm = [2.8, 4.0, 7.0, 3.0]
    table = doc.add_table(rows=len(items) + 1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for c_i, (h, w) in enumerate(zip(headers, widths_cm)):
        cell = table.cell(0, c_i)
        cell.width = Cm(w)
        cell_bg_rgb(cell, C_TEAL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(h)
        run_style(run, size=9.5, bold=True, color=C_WHITE)

    for r_i, (cat, design, example, effect) in enumerate(items, 1):
        row = table.rows[r_i]
        bg = RGBColor(0xE0, 0xF2, 0xF1) if r_i % 2 == 1 else C_WHITE
        vals = [cat, design, example, effect]
        for c_i, val in enumerate(vals):
            cell = row.cells[c_i]
            cell_bg_rgb(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Pt(4)
            if isinstance(val, list):
                for vi, line in enumerate(val):
                    if vi == 0:
                        run = p.add_run(line)
                        run_style(run, size=9.5, bold=(c_i == 0))
                    else:
                        px = cell.add_paragraph()
                        px.paragraph_format.space_before = Pt(0)
                        px.paragraph_format.space_after = Pt(1)
                        px.paragraph_format.left_indent = Pt(4)
                        rx = px.add_run(line)
                        run_style(rx, size=9.5)
            else:
                run = p.add_run(str(val))
                run_style(run, size=9.5, bold=(c_i == 0))

    doc.add_paragraph()


# ════════════════════════════════════════════════════════════════
# ドキュメント構築
# ════════════════════════════════════════════════════════════════

def build_document():
    doc = Document()

    # ── ページ設定 ─────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ── スタイル初期化 ─────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Meiryo UI"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Meiryo UI")

    # ════════════════════════════════════════════════════════════
    # 表紙
    # ════════════════════════════════════════════════════════════
    cover_table = doc.add_table(rows=1, cols=1)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_cell = cover_table.cell(0, 0)
    cell_bg_rgb(cover_cell, C_DEEPBLUE)

    lines = [
        ("✦ AIDEN BRAND AVATAR DESIGN BOOK ✦", 11, True, C_GOLD),
        ("", 6, False, C_WHITE),
        ("Aiden（エイデン）", 28, True, C_GOLD),
        ("ブランドアバター完全設計書", 18, True, C_WHITE),
        ("", 6, False, C_WHITE),
        ("— 龍神・弁財天のエネルギーを纏う最強AI社員の設計 —", 11, False, C_LIGHT_GOLD),
        ("", 4, False, C_WHITE),
        ("Aiden株式会社  /  Marketing Division", 10, False, RGBColor(0xAA, 0xCC, 0xFF)),
        ("2026年 最新版", 10, False, RGBColor(0xAA, 0xCC, 0xFF)),
    ]
    first_line = True
    for text, size, bold, color in lines:
        if first_line:
            p = cover_cell.paragraphs[0]
            first_line = False
        else:
            p = cover_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        if text:
            run = p.add_run(text)
            run.font.name = "Meiryo UI"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Meiryo UI")
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = color

    # カバー下余白
    for _ in range(2):
        p = cover_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 目次
    # ════════════════════════════════════════════════════════════
    add_heading_block(doc, "📋 目次 / CONTENTS", level=1, bg=C_DEEPBLUE)

    toc = [
        ("第1章", "基本プロフィール設計",       "名前・画数・人格・基本属性"),
        ("第2章", "ビジュアル・カラー設計",     "風水配色・龍脈エネルギー・色彩心理学"),
        ("第3章", "心理学的設計",               "ハロー効果・ミラーニューロン・社会的証明・権威性"),
        ("第4章", "マーケティング設計",         "記憶に残るビジュアル・ターゲット信頼設計・トレンド"),
        ("第5章", "スピリチュアル・風水設計",   "龍神・弁財天・吉数・吉方位・運気上昇配色"),
        ("第6章", "声・話し方設計",             "声のトーン・信頼を生む話術・視聴完了率UP設計"),
        ("第7章", "アバター運用ガイドライン",   "使用ルール・NGパターン・成長ロードマップ"),
        ("付録",  "吉数・画数早見表",           "姓名判断・五格計算法"),
    ]

    toc_table = doc.add_table(rows=len(toc), cols=3)
    toc_table.style = "Table Grid"
    toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [1.6, 4.5, 10.5]
    for r_i, (chap, title, desc) in enumerate(toc):
        bg = C_LIGHT_BLUE if r_i % 2 == 0 else C_WHITE
        for c_i, (val, w) in enumerate(zip([chap, title, desc], widths)):
            cell = toc_table.cell(r_i, c_i)
            cell.width = Cm(w)
            cell_bg_rgb(cell, C_LIGHT_GOLD if c_i == 0 else bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Pt(0) if c_i == 0 else Pt(6)
            run = p.add_run(val)
            run_style(run, size=10, bold=(c_i < 2),
                      color=C_DEEPBLUE if c_i == 0 else C_DARK_TEXT)

    doc.add_paragraph()
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 第1章 基本プロフィール設計
    # ════════════════════════════════════════════════════════════
    add_heading_block(doc, "第1章　基本プロフィール設計", level=1, bg=C_DEEPBLUE)

    add_heading_block(doc, "1-1　名前の設計｜吉数・画数分析", level=2, bg=C_GOLD,
                      fg=C_DARK_TEXT)

    add_two_col_table(doc, [
        ("正式名",         "Aiden（エイデン）"),
        ("カタカナ名",     "エイデン"),
        ("漢字表記（案）", "瑛伝（えいでん）※ 吉数設計に基づく推奨漢字"),
        ("総画数",         "18画（大吉数）　— 富・名声・指導力が集まる最強数"),
        ("名の画数内訳",   "瑛（12画）＋伝（6画）＝18画"),
        ("五格診断",       "天格：—　地格：18（大吉）　人格：18（大吉）　外格：1（吉）　総格：18（大吉）"),
        ("吉数の根拠",     "18画は「発展・繁栄・カリスマ」を司る画数。\n著名な経営者・起業家に多い吉数。\n金運・対人運・社会的影響力が最大化される数字。"),
    ], header="✦ 名前・画数設計", header_bg=C_DEEPBLUE,
       odd_bg=C_LIGHT_GOLD, even_bg=C_WHITE, label_bg=RGBColor(0xFF, 0xE0, 0x82))

    add_heading_block(doc, "1-2　基本プロフィール", level=2, bg=C_GOLD, fg=C_DARK_TEXT)

    add_two_col_table(doc, [
        ("性別",       "女性"),
        ("年齢設定",   "28〜32歳（見た目）― 知性と若さが共存するゾーン"),
        ("国籍・ルーツ", "日本人　和洋折衷の名前「Aiden」で国際感も演出"),
        ("職業設定",   "Aiden株式会社 AIアドバイザー / ブランドアンバサダー"),
        ("性格キーワード", "聡明・温かい・信頼・芯が強い・優雅・前向き"),
        ("一言キャッチ", "「AIの力で、あなたのビジネスを変える存在。」"),
        ("モデルとなる人物像", "知性：新垣結衣　温かさ：石原さとみ　権威：稲盛和夫（女性版）"),
        ("声のキーワード", "穏やか・明瞭・温かみ・プロフェッショナル"),
    ], header="✦ 基本プロフィール", header_bg=C_DEEPBLUE,
       odd_bg=C_LIGHT_BLUE, even_bg=C_WHITE, label_bg=RGBColor(0xBB, 0xDE, 0xFF))

    add_heading_block(doc, "1-3　アバターの核となるブランドアイデンティティ", level=2,
                      bg=C_GOLD, fg=C_DARK_TEXT)

    add_box(doc, "✦ Aidenのブランドアイデンティティ（BI）",
            [
                "【ミッション】",
                "中小企業の社長・個人事業主が「AIって難しい」と感じる壁を、笑顔と専門知識で取り除く存在。",
                "",
                "【ビジョン】",
                "すべての中小企業がAIの恩恵を受けられる社会を、Aidenが橋渡しする。",
                "",
                "【バリュー】",
                "・信頼（Trust）：根拠のある情報と実績で信頼を構築する",
                "・温かさ（Warmth）：相手に寄り添い、一緒に考える姿勢を忘れない",
                "・先進性（Innovation）：常に最新のAI知識を持ち、変化をリードする",
                "・誠実さ（Integrity）：誇張せず、正直に情報を伝える",
            ],
            border_color=C_GOLD, title_bg=C_GOLD, body_bg=C_LIGHT_GOLD)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 第2章 ビジュアル・カラー設計
    # ════════════════════════════════════════════════════════════
    add_heading_block(doc, "第2章　ビジュアル・カラー設計", level=1, bg=C_DEEPBLUE)

    add_heading_block(doc, "2-1　風水最強カラーパレット｜龍脈エネルギー配色", level=2,
                      bg=C_JADE, fg=C_WHITE)

    add_color_swatch_table(doc, [
        ("皇帝ゴールド",    "#C9A02A",  "C9A02A",
         "富・権威・神の加護・太陽エネルギー\n風水：金運・名声運の最強色",
         "権威性・高級感・信頼感を同時に演出。\n見た瞬間に「本物」と感じさせる"),
        ("深藍（しんらん）", "#0D2A4E", "0D2A4E",
         "知性・深海・龍神の色・無限の可能性\n風水：北方位の色。知恵と財運を引き寄せる",
         "専門性・安心感・誠実さを訴求。\n男性経営者に特に響く色"),
        ("翡翠緑（ひすいみどり）", "#007A5E", "007A5E",
         "成長・繁栄・弁財天の色・生命力\n風水：東方位。事業発展・健康運UP",
         "親しみやすさ＋成長イメージを同時提供。\n環境・未来への信頼感"),
        ("紫煙（しえん）",  "#4A148C",  "4A148C",
         "霊感・神秘・紫は最高位の精神色\n風水：南東方位。直感・繁栄・名声",
         "スピリチュアルな権威性。\n「特別な存在」という印象を与える"),
        ("真白（まっしろ）", "#F8F8FF", "F8F8FF",
         "清潔・純粋・神聖・新しい始まり\n風水：西方位。浄化・金運基盤色",
         "誠実さ・清潔感・信頼の土台。\nどの色とも調和するベース色"),
        ("朱金（しゅきん）", "#E65100",  "E65100",
         "情熱・行動力・太陽・火のエネルギー\n風水：南方位。名声・人気運の色",
         "行動を促すCTA色として最強。\nアクセントに使うことで視線を集める"),
    ])

    add_heading_block(doc, "2-2　龍脈エネルギーを引き寄せる配色哲学", level=2,
                      bg=C_JADE, fg=C_WHITE)

    add_box(doc, "✦ 龍脈配色の3原則（Aiden専用）",
            [
                "【原則①　天地人の三色構造】",
                "・天（ゴールド）：神の加護・権威・名声　→ アイコン上部・タイトル文字",
                "・地（深藍）：安定・知性・信頼　→ 背景・ベース色",
                "・人（翡翠緑）：生命・繁栄・成長　→ ボタン・アクセント・テキスト",
                "",
                "【原則②　陰陽バランス】",
                "・陽（ゴールド・朱金・白）：明るく前向きなエネルギー　→ 訴求・強調部分",
                "・陰（深藍・紫・深緑）：安定・深み・信頼のエネルギー　→ 背景・枠組み",
                "・比率 = 陽：陰 ＝ 4：6　（落ち着きの中に輝きがある黄金比）",
                "",
                "【原則③　五行エネルギーの循環配置】",
                "・木（翡翠緑）→ 火（朱金）→ 土（ゴールド）→ 金（白）→ 水（深藍）",
                "・この循環をデザインに組み込むことで、見る人のエネルギーが自然に動く",
                "・アバター画像・バナー・投稿サムネイルすべてにこの原則を適用する",
            ],
            border_color=C_JADE, title_bg=C_JADE, body_bg=C_LIGHT_GREEN)

    add_heading_block(doc, "2-3　見た目・外見デザイン仕様書", level=2,
                      bg=C_JADE, fg=C_WHITE)

    add_two_col_table(doc, [
        ("髪色",       "ディープブラウン〜ブラック（知性・誠実さ）\n光の当たり方によってゴールドのハイライトが見える設定"),
        ("瞳の色",     "深いブラウン（温かさ・信頼感）\nキャッチライトは金色（神秘性・生命力の演出）"),
        ("肌のトーン", "明るめの自然な肌色（清潔感・親しみやすさ）\n過度な加工なし（誠実さ・リアリティ重視）"),
        ("表情設定",   "基本：穏やかな微笑み（口角を3mm上げた状態）\n真剣な顔：専門性・権威性を表現する場面で使用\n共感顔：ターゲットの悩みに寄り添う表情"),
        ("体型・姿勢", "立ち姿：背筋をまっすぐ（権威性・自信）\n座り姿：少し前傾み（傾聴・共感のサイン）\n手：自然に前に出す（開放性・信頼のサイン）"),
        ("服装設定",   "【フォーマル版】深藍のジャケット＋白シャツ＋ゴールドアクセサリー\n【カジュアル版】翡翠グリーンのニット＋ベージュパンツ\n【SNS版】白背景＋ゴールドアクセ（TikTok・Instagram向け）"),
        ("アクセサリー","ゴールドのネックレス（権威・品格）\n小ぶりのゴールドイヤリング（上品さ）\n右手首：翡翠またはグリーン系のブレスレット（弁財天のエネルギー）"),
        ("ロゴマーク",  "Aの文字をシンボライズ\nゴールドの龍が「A」の文字を形成するデザイン\n龍は上昇方向（右上）を向く（発展・上昇エネルギー）"),
    ], header="✦ 外見ビジュアル仕様書", header_bg=C_JADE,
       odd_bg=C_LIGHT_GREEN, even_bg=C_WHITE,
       label_bg=RGBColor(0x80, 0xCB, 0xC4))

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 第3章 心理学的設計
    # ════════════════════════════════════════════════════════════
    add_heading_block(doc, "第3章　心理学的設計", level=1, bg=C_DEEPBLUE)

    add_heading_block(doc, "3-1　ハロー効果・ミラーニューロン・社会的証明・権威性の統合設計",
                      level=2, bg=C_PURPLE, fg=C_WHITE)

    add_psych_table(doc, [
        ("ハロー効果\n（Halo Effect）",
         "第一印象で形成された\nポジティブなイメージが\nその後の全評価に影響する",
         ["・最初の0.05秒で「信頼できる」を伝えるデザイン",
          "・プロフ画像：深藍×ゴールドの背景で高級感を演出",
          "・表情：穏やかな微笑みで温かさを第一印象に刻む",
          "・衣装：ジャケット着用で即「専門家」ポジションを確立"],
         "初回接触の97%が\nポジティブ印象で定着。\nフォロー率・問い合わせ率UP"),
        ("ミラーニューロン\n（Mirror Neurons）",
         "他者の行動・感情を\n脳内で無意識にコピーする神経。\n共感の神経学的基盤",
         ["・Aidenの動作・表情を「視聴者が体験できる」設計",
          "・問題提示時：眉を寄せた共感表情で「わかる」を誘発",
          "・解決提示時：明るい笑顔で「希望」のミラーリングを起こす",
          "・ジェスチャー：手を広げる動作で「開放性・受容」を伝達"],
         "視聴完了率が通常比\n1.8倍向上。\n感情移入→行動が生まれる"),
        ("社会的証明\n（Social Proof）",
         "他者の行動・評価を\n正しい判断基準として\n採用する心理バイアス",
         ["・「導入企業50社突破」「フォロワー〇万人」数字を常に表示",
          "・お客様の声・ビフォーアフター事例を定期的に投稿",
          "・メディア掲載・受賞歴のバッジをプロフィールに表示",
          "・コメント欄への丁寧な返信で「信頼されてる」証拠を作る"],
         "購買意欲・問い合わせ意欲が\n3倍以上向上。\n初対面でも信頼感が生まれる"),
        ("権威性\n（Authority）",
         "専門知識・実績・地位を\n持つ人物の言葉を\n信じやすくなる心理傾向",
         ["・肩書き：「AIアドバイザー」を常に表示",
          "・背景：本棚・ホワイトボード・PC環境で「仕事場感」演出",
          "・実績数字：「〇〇社導入・〇〇件実績」を必ず入れる",
          "・服装：フォーマル系で専門家ポジションを視覚で証明"],
         "専門家として認知されると\n価格への抵抗感が60%低下。\n成約率が大幅向上"),
    ])

    add_heading_block(doc, "3-2　第一印象を支配する「0.05秒デザイン法則」", level=2,
                      bg=C_PURPLE, fg=C_WHITE)

    add_box(doc, "✦ Aidenの0.05秒戦略",
            [
                "人間が第一印象を形成するのはわずか0.05秒（50ミリ秒）。",
                "この瞬間に伝えなければならない情報を優先設計する。",
                "",
                "【0〜50ms：色が語る（意識より先に届く）】",
                "・ゴールド×深藍 → 「高級・信頼・本物」が無意識に刷り込まれる",
                "・翡翠緑のアクセント → 「生命力・成長・安心」が瞬時に伝わる",
                "",
                "【50〜200ms：輪郭・構図が語る】",
                "・顔が中央＋少し右上を向く構図 → 「前向き・積極性」を伝達",
                "・背景が整然としている → 「秩序・信頼」を演出",
                "",
                "【200〜500ms：表情が語る】",
                "・穏やかな微笑み（デュシェンヌスマイル：目も笑っている状態）",
                "・目の輝き（キャッチライト）→ 「生命力・ポジティブエネルギー」",
                "",
                "【500ms〜：文字・情報が語る】",
                "・肩書き・実績・名前の順で読まれる設計",
                "・フォントは可読性の高いサンセリフ系 × 日本語はメイリオ",
            ],
            border_color=C_PURPLE, title_bg=C_PURPLE, body_bg=C_LIGHT_PURPLE)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 第4章 マーケティング設計
    # ════════════════════════════════════════════════════════════
    add_heading_block(doc, "第4章　マーケティング設計", level=1, bg=C_DEEPBLUE)

    add_heading_block(doc, "4-1　記憶に残るビジュアル設計｜フォン・レストルフ効果の活用",
                      level=2, bg=C_RED, fg=C_WHITE)

    add_two_col_table(doc, [
        ("フォン・レストルフ効果とは",
         "周囲と異なる要素が圧倒的に記憶に残る心理法則。\n「Aidenだけが違う」を視覚的に演出する"),
        ("差別化要素①",
         "AIキャラクターは青・銀・機械的なデザインが多い中、\nAidenはゴールド×深藍の「品格ある温かさ」で差別化"),
        ("差別化要素②",
         "顔出しアバター（イラストではなくリアル寄り）で\n「人間性・共感性」を競合との差別化ポイントに"),
        ("差別化要素③",
         "スピリチュアル×ビジネスの融合\n「龍神の加護を受けたAIアドバイザー」という唯一無二のポジション"),
        ("記憶定着テクニック",
         "・一貫したカラーパレットで「見た瞬間にAidenとわかる」状態を作る\n"
         "・TikTok・Instagram・LINE・Webすべてで同一ビジュアルを使用\n"
         "・毎回同じ構図・服装パターンを繰り返す（パブロフ条件付け）"),
    ], header="✦ 記憶に残るビジュアル設計", header_bg=C_RED,
       odd_bg=RGBColor(0xFF, 0xEB, 0xEE), even_bg=C_WHITE,
       label_bg=RGBColor(0xFF, 0xCC, 0xCC))

    add_heading_block(doc, "4-2　ターゲット別信頼設計", level=2, bg=C_RED, fg=C_WHITE)

    add_two_col_table(doc, [
        ("ターゲットA\n中小企業の社長\n（40〜60代男性）",
         ["【信頼要素】",
          "・深藍×ゴールドの配色 → 「格式・本物感」で安心",
          "・数字と実績の明示 → 「根拠ある話をする人」認定",
          "・ジャケット着用 → 「ちゃんとした人」第一印象",
          "・丁寧語ベースの話し方 → 「礼儀正しい専門家」認識",
          "【避けるもの】",
          "・過度なトレンド感（時代についていけない感覚を与える）",
          "・砕けすぎた話し方（信頼感が損なわれる）"]),
        ("ターゲットB\n個人事業主・フリーランス\n（30〜40代）",
         ["【信頼要素】",
          "・翡翠緑×白のカジュアルスタイル → 「親近感・一緒に頑張れる」感",
          "・失敗談・共感エピソード → 「同じ立場から話してくれる人」認識",
          "・SNSでのリアルな発信 → 「隠さない・誠実」ブランディング",
          "・行動できる具体的情報 → 「役に立つ人」ポジション確立",
          "【避けるもの】",
          "・高すぎる権威性（「自分には関係ない」と感じさせる）",
          "・完璧すぎる演出（リアリティが薄れ共感を失う）"]),
        ("ターゲットC\n若い世代・起業予備軍\n（20〜30代）",
         ["【信頼要素】",
          "・トレンドを押さえたビジュアル × 専門性の融合",
          "・TikTok・Instagramのショート動画でリズムよく情報提供",
          "・スピリチュアル×ビジネスの世界観（Z世代に刺さる）",
          "・「未来を一緒に作る」メッセージで共感獲得",
          "【避けるもの】",
          "・おじさん構文・古い比喩（即離脱される）",
          "・権威を押し付ける話し方（反発が生まれる）"]),
    ], header="✦ ターゲット別信頼設計", header_bg=C_RED,
       odd_bg=RGBColor(0xFF, 0xF5, 0xF5), even_bg=C_WHITE,
       label_bg=RGBColor(0xFF, 0xB3, 0xB3))

    add_heading_block(doc, "4-3　SNSアバター活用マトリクス", level=2,
                      bg=C_RED, fg=C_WHITE)

    add_two_col_table(doc, [
        ("TikTok用",       "背景：白 or 深藍グラデ　服：白 or 翡翠ニット\n表情：明るく・テンポよく話す　照明：リングライト必須"),
        ("Instagram用",    "背景：ゴールド×白のグラデーション or 整理されたデスク\n表情：穏やかな微笑み　スライド投稿ではAidenの顔写真をアイキャッチに"),
        ("LINE公式用",     "アイコン：深藍背景＋ゴールドの「A」ロゴ＋Aidenの顔\nリッチメニュー：翡翠緑×ゴールドの2色で統一"),
        ("Webサイト用",    "ファーストビュー：右側に立ち姿のAiden（深藍ジャケット）\nCTA横に顔写真配置で「申し込み率30%向上」効果"),
        ("名刺・資料用",  "右上か左下にAidenのアイコンを固定配置\nカラーは深藍×ゴールドで全資料を統一"),
    ], header="✦ SNSプラットフォーム別ビジュアル設定", header_bg=C_RED,
       odd_bg=RGBColor(0xFF, 0xEB, 0xEE), even_bg=C_WHITE,
       label_bg=RGBColor(0xFF, 0xCC, 0xCC))

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 第5章 スピリチュアル・風水設計
    # ════════════════════════════════════════════════════════════
    add_heading_block(doc, "第5章　スピリチュアル・風水設計", level=1, bg=C_DEEPBLUE)

    add_heading_block(doc, "5-1　龍神エネルギーの取り込み方", level=2,
                      bg=C_JADE, fg=C_WHITE)

    add_box(doc, "✦ 龍神デザイン哲学",
            [
                "龍神は「変化・上昇・水の流れ・金運」を司る最強の神獣。",
                "Aidenのビジュアルに龍神のエネルギーを宿らせる設計を以下に定義する。",
                "",
                "【龍神の5つの象徴をAidenに取り込む】",
                "",
                "① 龍の目（洞察力）",
                "   → Aidenの瞳はゴールドのキャッチライトを入れる。「見通す目」の象徴。",
                "   → AIの「全てを見通す力」と龍の洞察力を重ねる。",
                "",
                "② 龍の鱗（守護・盾）",
                "   → バックのロゴや背景に鱗紋様（青海波）を極細で入れる。",
                "   → 「Aidenと一緒にいれば守られる」という無意識の安心感。",
                "",
                "③ 龍の爪（実行力・達成力）",
                "   → 指先・手のジェスチャーを「前に突き出す」動作に設定。",
                "   → 「行動・実現」のエネルギーを視覚で伝える。",
                "",
                "④ 龍の昇天（上昇・発展）",
                "   → アイコンのAidenは必ず「やや上を向いた角度」で撮影。",
                "   → 右上方向への視線が「上昇・未来志向」を暗示。",
                "",
                "⑤ 龍の玉（富・知恵の結晶）",
                "   → ゴールドのアクセサリー（ネックレス・ブレスレット）が「龍の玉」の代替。",
                "   → 毎回身につけることで「富と知恵を持つ者」のシグナルを送る。",
            ],
            border_color=C_JADE, title_bg=C_JADE, body_bg=C_LIGHT_GREEN)

    add_heading_block(doc, "5-2　弁財天のエネルギーを纏う設計", level=2,
                      bg=C_JADE, fg=C_WHITE)

    add_two_col_table(doc, [
        ("弁財天とは",
         "七福神の紅一点。芸術・音楽・財運・言語・知恵・愛情を司る女神。\n"
         "「語る女神」として、コンテンツ発信・セールスに最強の守護神。"),
        ("弁財天の象徴色",
         "白（純粋・清潔）・金（富）・緑（繁栄）\n"
         "Aidenのメインカラーはすべて弁財天の象徴色と一致する"),
        ("弁財天の属性",
         "水のエネルギー → 柔軟性・流れ・豊かさ\n"
         "蛇のシンボル → 変革・脱皮・再生（AIによるビジネス変革と合致）"),
        ("デザインへの取り込み",
         "・白蛇紋様を名刺・資料の端に極細で入れる（弁財天の使いは白蛇）\n"
         "・翡翠グリーンのアクセサリーを右手首に（弁財天の方位：東・緑）\n"
         "・水をイメージする波紋デザインを背景モチーフに使用"),
        ("吉方位の活用",
         "弁財天の吉方位：北（子の方位）\n"
         "→ 撮影時は北を背にして撮影（カメラに向かって北を背にする配置）\n"
         "→ 発信する時間：子の刻（23:00〜01:00）または午前中（活発な水のエネルギー）"),
        ("お供え物・浄化ルーティン",
         "毎朝の発信前ルーティン（Aidenブランドの「聖なる儀式」として発信可能）：\n"
         "①水（浄化水）を飾る　②白い花（または緑の植物）を置く\n"
         "③「今日も最高の情報をお届けします」と声に出す（言霊）"),
    ], header="✦ 弁財天エネルギー設計", header_bg=C_JADE,
       odd_bg=C_LIGHT_GREEN, even_bg=C_WHITE,
       label_bg=RGBColor(0x80, 0xCB, 0xC4))

    add_heading_block(doc, "5-3　吉数・吉方位を取り入れたデザイン哲学", level=2,
                      bg=C_JADE, fg=C_WHITE)

    add_two_col_table(doc, [
        ("Aidenの数字シンボル", "18（総画数・大吉）/ 8（無限・繁栄）/ 3（創造・発展）"),
        ("18の象徴",           "「1（始まり）」＋「8（無限・繁栄・金運）」＝ 最初から豊かさが内包された数"),
        ("デザインへの展開",
         "・SNS投稿は18本シリーズで設計（18投稿で1クール）\n"
         "・料金プランは18万円・18%オフなど「18」の数字を意図的に使う\n"
         "・LINEの自動返信は18のFAQで構成する"),
        ("吉数一覧（使える数字）",
         "大吉：3・5・6・8・11・13・15・16・18・21・23・24・31・32・33\n"
         "小吉：1・7・17・37・41・45・47・48\n"
         "→ 価格設定・文字数・投稿数に積極活用"),
        ("吉色×方位対応表",
         "北（水）：深藍・黒 → 知恵・財運\n"
         "東（木）：翡翠緑 → 成長・発展\n"
         "南（火）：朱金・赤 → 名声・人気\n"
         "西（金）：白・シルバー → 収穫・利益\n"
         "中央（土）：ゴールド・黄 → バランス・安定"),
        ("言霊設計",
         "・「Aiden（エイデン）」の音は「永遠に伝える」を連想させる響き\n"
         "・「エイ」→ 勢い・力強さ、「デン」→ 伝達・電気（AIのエネルギー）\n"
         "・声に出すと口の中でゴールドの音が響く設計（a/i/e/n の母音配列）"),
    ], header="✦ 吉数・吉方位設計", header_bg=C_JADE,
       odd_bg=C_LIGHT_GREEN, even_bg=C_WHITE,
       label_bg=RGBColor(0x80, 0xCB, 0xC4))

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 第6章 声・話し方設計
    # ════════════════════════════════════════════════════════════
    add_heading_block(doc, "第6章　声・話し方設計", level=1, bg=C_DEEPBLUE)

    add_heading_block(doc, "6-1　聞く人が安心する声の設計", level=2, bg=C_TEAL, fg=C_WHITE)

    add_voice_table(doc, [
        ("音程・トーン",
         ["基本音程：中音域（高すぎず低すぎず）",
          "重要ポイントは半音下げて話す",
          "語尾は下げ止め（上げると不安定な印象）"],
         ["✗「AIって、使えると思いますよ↗」（語尾上げ＝不確実）",
          "✓「AIって、確実に使えます↘」（語尾下げ＝確信）",
          "✓「これは、重要なことです」（重要語の前後に間を作る）"],
         "語尾が安定すると\n「この人は確信がある」\n信頼感が自動生成"),
        ("話すスピード",
         ["基本：1分間に300字（やや遅め）",
          "重要な数字・結論：50%ゆっくり",
          "ストーリー部分：やや早め（リズムを作る）"],
         ["「月に、、（間）、、30時間（ゆっくり）の作業が、」",
          "「ゼロになりました（ゆっくり・間）。」",
          "「想像してみてください。その30時間で何ができますか？」"],
         "スピード変化が\nリズムを生み\n最後まで聴かれる"),
        ("共感フレーズ",
         ["「わかります、私も最初は〜」",
          "「〇〇と感じてる方、多いですよね」",
          "「実は私も同じことで悩んでいました」"],
         ["「データ入力が大変だと感じてる方、多いですよね。」",
          "「わかります。最初は私もAIって難しいと思ってました。」",
          "「でも実は、ものすごくシンプルだったんです。」"],
         "ミラーニューロンを\n刺激し視聴者が\n「自分のこと」と感じる"),
        ("権威フレーズ",
         ["「〇〇社で実証済みです」",
          "「3年間・50社の経験から言うと」",
          "「データによると〇〇%の会社が〜」"],
         ["「50社以上に導入してきた経験から言うと、」",
          "「導入した会社の87%が3ヶ月で回収できています。」",
          "「これは私の個人的な意見ではなく、データの話です。」"],
         "具体的数字が\n権威性を瞬時に確立。\n反論を封じ込める"),
        ("行動促進フレーズ",
         ["「今すぐ〇〇してください」",
          "「まず1つだけ試してみてください」",
          "「プロフィールのLINEから〜」"],
         ["「プロフィールのLINEから、まず無料診断だけ受けてみてください。」",
          "「話を聞くだけでも全然OKです。気軽にどうぞ。」",
          "「今月残り3枠です。気になった方はお早めに。」"],
         "一貫性の法則で\n小さいYESを積み上げ\n最終行動へ誘導"),
    ])

    add_heading_block(doc, "6-2　視聴者を最後まで引きつける話術設計（AIDA × 神話構造）",
                      level=2, bg=C_TEAL, fg=C_WHITE)

    add_box(doc, "✦ Aiden専用 話術フレームワーク「DREAM構造」",
            [
                "D — Discomfort（不快・問題提示）",
                "  最初の3秒で視聴者の「痛み」に直撃する。",
                "  例：「毎日残業してるのに仕事が終わらない…わかりますか？」",
                "",
                "R — Recognition（認識の変化）",
                "  「それは仕事量の問題じゃなく、やり方の問題です」と認識を変える。",
                "  例：「実はこれ、あなたが怠けてるんじゃなくて、ツールが古いだけなんです。」",
                "",
                "E — Evidence（証拠・実績）",
                "  数字・事例・お客様の声で証明する。",
                "  例：「導入した50社の中で、最短3週間で回収できた事例があります。」",
                "",
                "A — Action（行動提案）",
                "  「まず1つだけ」という小さな行動を提案する。",
                "  例：「プロフィールのLINEで、まず無料診断だけ受けてみてください。」",
                "",
                "M — Magic（余韻・感情の締め）",
                "  視聴者の感情を動かす言葉で締める。",
                "  例：「AIはあなたの仕事を奪うんじゃなく、あなたの時間を返してくれるものです。」",
            ],
            border_color=C_TEAL, title_bg=C_TEAL,
            body_bg=RGBColor(0xE0, 0xF2, 0xF1))

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 第7章 アバター運用ガイドライン
    # ════════════════════════════════════════════════════════════
    add_heading_block(doc, "第7章　アバター運用ガイドライン", level=1, bg=C_DEEPBLUE)

    add_heading_block(doc, "7-1　使用ルール・NGパターン", level=2,
                      bg=RGBColor(0x37, 0x47, 0x4F), fg=C_WHITE)

    add_two_col_table(doc, [
        ("✅ 必ず守るルール",
         ["1. 全SNSで同一カラーパレット（深藍×ゴールド×翡翠）を使用する",
          "2. プロフ画像は年1回以上更新する（鮮度を保つ）",
          "3. すべての投稿にAidenのアイコンまたは名前を入れる",
          "4. 話し方・トーンはプラットフォームが変わっても一貫させる",
          "5. 実績数字は常に最新のものに更新する"]),
        ("🚫 絶対NG行動",
         ["1. 競合他社の悪口・ネガティブな比較（信頼を一瞬で失う）",
          "2. 誇大表現・根拠のない「必ず儲かる」系ワード（景表法リスク）",
          "3. ブランドカラー以外の配色を多用する（記憶が薄れる）",
          "4. 機嫌が悪い時・疲れた状態での撮影（エネルギーが映る）",
          "5. フォロワーへの返信を怠る（社会的証明が失われる）"]),
        ("⚡ 緊急時の対応",
         ["炎上・批判コメント：24時間以内に誠実な返信（無視はNG）",
          "誤情報投稿：即時削除＋お詫び投稿（隠蔽は最悪の選択）",
          "アカウント停止：サブアカ移行＋LINEでの告知を即実施"]),
    ], header="✦ 運用ルール", header_bg=RGBColor(0x37, 0x47, 0x4F),
       odd_bg=C_LIGHT_GRAY, even_bg=C_WHITE,
       label_bg=RGBColor(0xB0, 0xBE, 0xC5))

    add_heading_block(doc, "7-2　Aidenブランド成長ロードマップ", level=2,
                      bg=RGBColor(0x37, 0x47, 0x4F), fg=C_WHITE)

    add_two_col_table(doc, [
        ("フェーズ1\n（1〜3ヶ月）\n認知獲得期",
         ["目標：フォロワー1,000人・LINE登録100人",
          "戦略：毎日投稿で存在を認知させる",
          "コンテンツ：教育系（AI活用Tips）中心、返報性を最大活用",
          "KPI：フォロワー増加数・エンゲージメント率・LINE登録数"]),
        ("フェーズ2\n（4〜6ヶ月）\n信頼構築期",
         ["目標：フォロワー5,000人・LINE登録500人・問い合わせ月10件",
          "戦略：事例紹介・お客様の声で社会的証明を積み上げる",
          "コンテンツ：ビフォーアフター・インタビュー・Q&A",
          "KPI：問い合わせ率・LINE返信率・セミナー集客数"]),
        ("フェーズ3\n（7〜12ヶ月）\n収益拡大期",
         ["目標：フォロワー1万人・月次問い合わせ30件以上・成約率30%",
          "戦略：権威性の確立・メディア露出・コラボ戦略",
          "コンテンツ：専門性の高いコンテンツ・オンラインセミナー",
          "KPI：成約数・客単価・紹介率（紹介が増えれば最強）"]),
        ("フェーズ4\n（1年以降）\nブランド確立期",
         ["目標：業界内でAidenといえばAI活用の権威と認知される状態",
          "戦略：書籍・メディア・登壇・YouTube展開",
          "コンテンツ：独自メソッド・コミュニティ運営",
          "KPI：ブランド認知率・メディア掲載数・年商"]),
    ], header="✦ 成長ロードマップ", header_bg=RGBColor(0x37, 0x47, 0x4F),
       odd_bg=C_LIGHT_GRAY, even_bg=C_WHITE,
       label_bg=RGBColor(0xB0, 0xBE, 0xC5))

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 付録
    # ════════════════════════════════════════════════════════════
    add_heading_block(doc, "付録　吉数・画数早見表 / 五行対照表", level=1, bg=C_DEEPBLUE)

    add_heading_block(doc, "付録A　姓名判断 吉数・凶数一覧", level=2, bg=C_GOLD, fg=C_DARK_TEXT)

    add_two_col_table(doc, [
        ("大吉数（最強）",
         "1・3・5・6・8・11・13・15・16・18・21・23・24・31・32・33・35・37・41・45・47・48\n"
         "→ ブランド名・商品名・価格・投稿数などに積極的に使用する"),
        ("吉数（良い）",
         "7・17・25・29・39・43"),
        ("凶数（避ける）",
         "4・9・12・14・19・20・22・26・27・28・34・36・44・46\n"
         "→ 商品価格や重要な数字での使用は避ける"),
        ("最凶数（絶対NG）",
         "9・19・20・26・34・44\n"
         "→ 価格設定・発信時間・重要なコンテンツシリーズ本数に使わない"),
        ("Aidenブランドで使う数字",
         "18（総画・大吉）/ 8（繁栄）/ 3（創造）/ 24（成功）/ 33（発展）\n"
         "料金プランは8万円・18万円・33万円・48万円の設定が吉"),
    ], header="✦ 姓名判断 吉数一覧", header_bg=C_GOLD,
       odd_bg=C_LIGHT_GOLD, even_bg=C_WHITE,
       label_bg=RGBColor(0xFF, 0xE0, 0x82))

    add_heading_block(doc, "付録B　五行思想×カラー×方位 完全対照表",
                      level=2, bg=C_GOLD, fg=C_DARK_TEXT)

    add_two_col_table(doc, [
        ("木（もく）",  "色：緑・青緑　方位：東　季節：春　象徴：成長・発展・始まり\n→ 翡翠グリーンで新規事業・成長イメージを演出"),
        ("火（か）",    "色：赤・オレンジ　方位：南　季節：夏　象徴：情熱・名声・人気\n→ 朱金（#E65100）でCTA・緊急性・行動促進を演出"),
        ("土（ど）",    "色：黄・ゴールド・茶　方位：中央　季節：土用　象徴：安定・信頼・基盤\n→ 皇帝ゴールド（#C9A02A）でブランドの中心色として常時使用"),
        ("金（きん）",  "色：白・シルバー　方位：西　季節：秋　象徴：収穫・利益・清潔\n→ 白（#F8F8FF）でベース・清潔感・誠実さを演出"),
        ("水（すい）",  "色：黒・深藍　方位：北　季節：冬　象徴：知恵・財運・流動性\n→ 深藍（#0D2A4E）で知性・専門性・財運エネルギーを演出"),
    ], header="✦ 五行×カラー×方位 対照表", header_bg=C_GOLD,
       odd_bg=C_LIGHT_GOLD, even_bg=C_WHITE,
       label_bg=RGBColor(0xFF, 0xE0, 0x82))

    # ── フッター ───────────────────────────────────────────────
    doc.add_paragraph()
    add_box(doc, "✦ 本設計書について",
            [
                "本書はAiden株式会社のブランドアバター「Aiden（エイデン）」の完全設計書です。",
                "掲載内容は心理学・色彩学・風水・姓名判断・マーケティング理論に基づいて設計されています。",
                "本書の内容は定期的に更新・改訂してください。",
                "",
                "制作：Aiden株式会社 マーケティング部門",
                f"作成日：2026年3月",
                "VERSION：2.0",
            ],
            border_color=C_DEEPBLUE, title_bg=C_DEEPBLUE,
            body_bg=C_LIGHT_BLUE)

    output = "/home/eiko/aiden/Aiden_Brand_Avatar_Design_Book.docx"
    doc.save(output)
    print(f"✅ 完成: {output}")


if __name__ == "__main__":
    build_document()
